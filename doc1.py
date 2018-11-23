from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import sim
import numpy as np
                                                 
import time
from leer_preguntas import crear_preguntas
import qr

from sim import ecuacion_to_jpg

def crear_docx(numero=1):
	respuestas=""
	
	preguntas = crear_preguntas()
	for pregunta in preguntas:
		#print(pregunta.respuesta)
		respuestas = respuestas+str(1+pregunta.respuestas.index(pregunta.respuesta))

	print(respuestas)

	qr.crear_qr(respuestas)

	document = Document()

	from docx.shared import Pt

	style = document.styles['Normal']
	paragraph_format = style.paragraph_format
	paragraph_format.space_before = Pt(0)
	paragraph_format.space_after = Pt(0)
	font = style.font
	font.name = 'Arial'
	font.size = Pt(13)
	style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


	document.add_paragraph().add_run('UNIDAD EDUCATIVA "LA ALBORADA"').bold=True
	last_paragraph = document.paragraphs[-1] 
	last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
	last_paragraph.style.font.size=Pt(18)


	p=document.add_paragraph()
	p.add_run('Nombre:').bold=True
	p.add_run('.........................................\t')
	p.add_run('Paralelo:').bold=True
	p.add_run('.........................')

	p=document.add_paragraph()
	p.add_run('Fecha:').bold=True
	p.add_run('.....................................')


	p=document.add_paragraph("* SELECCIONE EN LA TABLA A CONTINUACION, EL LITERAL CORRECTO")
	p=document.add_paragraph("* NO SE ACEPTARAN RESPUESTAS VARIAS")
	p=document.add_paragraph("* NO TACHONES EN LA TABLA")
	#last_paragraph = document.paragraphs[-1] 
	#last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

	p=document.add_paragraph()
	table = document.add_table(rows=1, cols=5, style = 'Table Grid')
	#table.style = 'TableGrid'
	hdr_cells = table.rows[0].cells
	hdr_cells[0].paragraphs[0].add_run().add_picture("qr.jpg",width=Inches(1.2))
	hdr_cells[1].text = 'A'
	hdr_cells[2].text = 'B'
	hdr_cells[3].text = 'C'
	hdr_cells[4].text = 'D'

	for i in range(1,len(preguntas)+1):
		row_cells = table.add_row().cells
		row_cells[0].text=str(i)

	last_paragraph = document.paragraphs[-1] 
	last_paragraph.style.font.size=Pt(13)
	last_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

	from lxml import etree

	i=2
	tam=1.2
	numero_preguntas = len(preguntas)
	print("NUM_PREG"+str(numero_preguntas))
	for pregunta in preguntas:
		p = document.add_paragraph("", style='List Number')
		p.add_run(pregunta.literal).bold=True
		if pregunta.ecuacion:
			if ".jpg" in pregunta.ecuacion or ".png" in pregunta.ecuacion:
				p=document.add_paragraph()
				p.add_run().add_picture(pregunta.ecuacion,width=Inches(1.7))
				
			else:
				preg=sim.ecuacion_to_latex(pregunta.ecuacion)
				
				if isinstance(preg, (np.ndarray, np.generic) ):
					sim.crear_foto(preg,"preg.png",esMatriz=True)
					tam=1.4
				else:#no es matriz... ecuacion imagino...
					sim.crear_foto(preg,"preg.png")
					if len(preg)>70:
						tam=3
						#print("1P")
						#print(preg)
					elif len(preg)>60:
						#print("2")
						#print(preg)
						tam=2.2
					elif len(preg)>45:
						tam=1.9
					elif len(preg)>30:
						tam=1.6
					elif len(preg)>30:
						tam=1.3
					else:
						tam=0.9
				p=document.add_paragraph()
				p.add_run().add_picture("preg.png",width=Inches(tam))
				#print("Pr:"+preg+repr(tam))
				last_paragraph = document.paragraphs[-1] 
				last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
				
			
		sol = 97
		if pregunta.respuestas:
			for resp in pregunta.respuestas:
			
				res=sim.ecuacion_to_latex(resp)
				p=document.add_paragraph(chr(sol)+") ")
				if len(res)>100:
						tam=2.7
						#print("r1")
						#print(res)
				elif len(res)>90:
						tam=2.3
						#print("r1")
						#print(res)
				elif len(res)>75:
						#print("r2")
						tam=2.0
						#print(res)
				elif len(res)>60:
						tam=1.7
						#print("r3")
						#print(res)
				elif len(res)>45:
						tam=1.4
						#print("r4")
						#print(res)
				elif len(res)>30:
						tam=1.1
						#print("r5")
						#print(res)
				elif len(res)>15:
						tam=0.8
						#print("r6")
						#print(res)
				else :
						tam=0.55
						#print("r7")
						#print(res)
				if isinstance(res, (np.ndarray, np.generic) ):
					sim.crear_foto(res,"respM.png",tam=120,esMatriz=True)
					#print("Matrix")
					p.add_run().add_picture("respM.png",width=Inches(tam))
				else:
					sim.crear_foto(res,"resp.png",tam=130)
					p.add_run().add_picture("resp.png",width=Inches(tam))
				
				sol=sol+1
		print("Pregunta: "+str(i))
		if i%2==0 and i<numero_preguntas:
			document.add_page_break()
			print("NEXT PAGE")
		else:
			if not i==numero_preguntas+1:
				for j in range(0,4):
					#print("*Enter*")
					p = document.add_paragraph()
					run = p.add_run()
					run.add_break()
				#run.add_break()
		i=i+1	

		
	nombre_documento = 'exam'+str(numero)+'.docx'
	document.save(nombre_documento)
	print(nombre_documento+" creado")
	import os
	try:
		os.remove("qr.jpg")
		os.remove("preg.png")
		os.remove("resp.png")
		os.remove("respM.png")
	except:
		pass
